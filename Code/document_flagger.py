"""
=============================================================================
Project: Automated Compliance Matrix Generator
Description: This script reads government contract PDFs and Word docs, 
             builds a color-coded spreadsheet, and outputs a highlighted 
             PDF/Word Doc showing exactly where each clause was found.
=============================================================================
"""

"""
AI Use Statement:   Gemini 3.1 Pro was used extensively to aid in cleaning up, debugging, and commenting on code.
                    The AI provided suggestions on code structure, error handling, and optimization, but all final decisions and 
                    implementations were made by the human author.
"""

import os               # Helps the script find folders and files on the computer
import re               # Helps find specific text patterns (like exact clause numbers)
import pandas as pd     # Used to read, build, and save Excel and CSV spreadsheets
import pdfplumber       # Used to read text from normal digital PDFs
import pytesseract      # Used to read text from scanned pictures/documents (OCR)
from pdf2image import convert_from_path  # Turns PDF pages into pictures if needed
import docx                                 # Used to read and modify Word Documents
from docx.enum.text import WD_COLOR_INDEX   # Used to pick the yellow highlight color for Word
import fitz                                 # Used to physically draw highlights on PDFs
from datetime import datetime            # Adds the current time to our output files
from openpyxl import load_workbook          # Used to open the saved Excel file for coloring
from openpyxl.styles import PatternFill     # Used to paint the cell background colors

# =============================================================================
# --- Folder Setup ---
# =============================================================================

# The script looks for this exact column name in the database files to match everything up.
CLAUSE_COL_NAME = 'Clause' 

# Find exactly where this code is saved on the computer so it works on any machine.
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Look one folder up to find the main project folder.
if os.path.basename(SCRIPT_DIR).lower() == 'code':
    PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
else:
    PROJECT_DIR = SCRIPT_DIR

# Set the paths for where our data lives and where to save the results.
DATABASE_DIR = os.path.join(PROJECT_DIR, 'Database')
SOLICITATIONS_DIR = os.path.join(PROJECT_DIR, 'Solicitations')
OUTPUT_DIR = os.path.join(PROJECT_DIR, 'Output')

# =============================================================================
# --- Project Functions ---
# =============================================================================

def clean_headers(columns):
    """
    Cleans up messy column names in the Excel files. 
    Sometimes the files have weird spaces or asterisks in the headers. 
    This fixes them so they all match perfectly.
    """
    clean_cols = []
    for col in columns:
        c = str(col).replace('\n', ' ').replace('*', '').strip()
        c = re.sub(r'\s+', ' ', c)
        clean_cols.append(c)
    return clean_cols

def load_databases(db_dir):
    """
    Reads all our separate agency files (FAR, DFARS, NASA, etc.) from the 
    Database folder and combines them into one massive master dictionary.
    """
    print(f"Loading database files from: {db_dir}...")
    all_dataframes = [] 
    
    if not os.path.exists(db_dir):
        print(f"Error: Database folder '{db_dir}' not found.")
        return None

    # Go through every file in the database folder
    for filename in os.listdir(db_dir):
        filepath = os.path.join(db_dir, filename)
        
        # Skip hidden files or the old Definitions matrix
        if filename.startswith('~') or filename.startswith('.') or 'Definitions' in filename or filename == 'Contract Ts&Cs Matrix.xlsm':
            continue 

        try:
            # Open the file depending on if it's Excel or CSV
            if filename.lower().endswith(('.xlsx', '.xlsm')):
                df = pd.read_excel(filepath, engine='openpyxl')
            elif filename.lower().endswith('.xls'):
                df = pd.read_excel(filepath)
            elif filename.lower().endswith('.csv'):
                try:
                    df = pd.read_csv(filepath, encoding='utf-8', on_bad_lines='skip')
                except UnicodeDecodeError:
                    df = pd.read_csv(filepath, encoding='latin1', on_bad_lines='skip')
            else:
                continue 
            
            # Clean up the column names
            df.columns = clean_headers(df.columns)
            
            # Filter out junk rows (like instructions or blank lines)
            if CLAUSE_COL_NAME in df.columns:
                df[CLAUSE_COL_NAME] = df[CLAUSE_COL_NAME].astype(str).str.strip()
                df = df[df[CLAUSE_COL_NAME].str.contains(r'\d', na=False)] # Must contain a number
                df = df[df[CLAUSE_COL_NAME].str.len() < 30] # Can't be a whole paragraph
                
                all_dataframes.append(df)
            else:
                # If the file doesn't have a 'Clause' column, warn the user
                print(f"  ! Skipped: {filename} - Missing '{CLAUSE_COL_NAME}' column.")
            
        except Exception as e:
            print(f"  ! Error loading {filename}: {e}")

    if not all_dataframes:
        return None

    # Combine everything together into one big list
    combined_df = pd.concat(all_dataframes, ignore_index=True)
    print(f"Successfully built a master database of {len(combined_df)} total clauses.")
    return combined_df

# --- Text Extraction ---

def extract_text_from_pdf(pdf_path):
    """
    Tries to read the PDF normally. If it realizes the PDF is just a scanned 
    picture, it takes extra steps to take pictures of the pages and read them.
    """
    text = ""
    print(f"Extracting text from PDF: {os.path.basename(pdf_path)}...")
    
    # Try reading the text the fast, normal way
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"  ! Error reading PDF {pdf_path}: {e}")
        return None

    # If we barely found any text, it's probably a scanned document.
    if len(text.strip()) < 50:
        print("  ! No text found. Looks like a scan. Reading images now...")
        try:
            images = convert_from_path(pdf_path)
            ocr_text = ""
            for i, image in enumerate(images):
                print(f"    - Reading page {i + 1} of {len(images)}...")
                ocr_text += pytesseract.image_to_string(image) + "\n"
            
            text = ocr_text
            if not text.strip():
                print("  ! Warning: Could not find any text in the scan either. Skipping.")
                return None
                
        except Exception as e:
            print(f"  ! Image reading failed: {e}")
            return None

    return text

def extract_text_from_docx(docx_path):
    """
    Reads the text out of a standard Word document, including any text
    hidden inside tables.
    """
    print(f"Extracting text from Word Document: {os.path.basename(docx_path)}...")
    text = ""
    try:
        doc = docx.Document(docx_path)
        # Read standard paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Dig into tables and read the cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"  ! Error reading Word document {docx_path}: {e}")
        return None

# --- Cross-Referencing & Formatting ---

def find_clauses_from_db(text, master_clauses):
    """
    Looks through the text we pulled from the contract and checks if any 
    of our database clauses are in it (like a massive CTRL+F).
    """
    found_clauses = []
    for clause in master_clauses:
        # We check for exact matches so "52.2" doesn't accidentally trigger "52.212-4"
        pattern = r'\b' + re.escape(clause) + r'\b'
        if re.search(pattern, text):
            found_clauses.append(clause)
    return sorted(found_clauses)

def generate_compliance_matrix(found_clauses, master_df):
    """
    Builds the final Excel spreadsheet by matching the clauses we found in the 
    contract with their full details (Title, Status, etc.) from our database.
    """
    if master_df is None:
        return None

    matrix_rows = []
    headers = master_df.columns.tolist()
    print(f"Cross-referencing {len(found_clauses)} clauses with the combined database...")

    # Grab the full row information for every clause we found
    for clause in found_clauses:
        matching_rows = master_df[master_df[CLAUSE_COL_NAME] == clause]
        for _, row in matching_rows.iterrows():
            matrix_rows.append(row.tolist())

    # Put it all into a clean spreadsheet format and sort it alphabetically
    matrix_df = pd.DataFrame(matrix_rows, columns=headers)
    matrix_df = matrix_df.sort_values(by=CLAUSE_COL_NAME).reset_index(drop=True)
    return matrix_df

def apply_color_coding(filepath):
    """
    Opens the completed spreadsheet and paints the cells Green, Yellow, or Red 
    based on the client's internal compliance rubric.
    """
    print("  - Applying rubric color-coding to the spreadsheet...")
    
    fill_ok = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')      # Green
    fill_c = PatternFill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')       # Yellow
    fill_remove = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')  # Red

    wb = load_workbook(filepath)
    ws = wb.active

    # Check every single cell in the spreadsheet
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            val = str(cell.value).strip().lower() if cell.value else ""
            if val == 'ok':
                cell.fill = fill_ok
            elif val == 'c':
                cell.fill = fill_c
            elif val == 'remove':
                cell.fill = fill_remove

    wb.save(filepath)

# --- Highlighting Functions ---

def highlight_pdf(input_path, output_path, found_clauses):
    """
    Opens the original PDF, finds the physical coordinates of the clauses we found,
    draws a yellow highlight box over them, and saves a new copy.
    """
    print("  - Generating highlighted PDF...")
    try:
        doc = fitz.open(input_path)
        for page in doc:
            for clause in found_clauses:
                # Find exactly where the text lives on the page
                text_instances = page.search_for(clause)
                for inst in text_instances:
                    # Draw a yellow box over it
                    highlight = page.add_highlight_annot(inst)
                    highlight.update() 
        doc.save(output_path, garbage=4, deflate=True, clean=True)
        doc.close()
        print(f"  -> Highlighted PDF saved to: {output_path}")
    except Exception as e:
        print(f"  ! Failed to highlight PDF: {e}")

def _apply_highlights_to_paragraph(paragraph, found_clauses):
    """
    Word documents are tricky. To highlight a specific word, we have to erase 
    the paragraph and redraw it piece-by-piece with the yellow highlight added in.
    """
    # Find which clauses are actually in this specific paragraph
    clauses_in_para = [c for c in found_clauses if re.search(r'\b' + re.escape(c) + r'\b', paragraph.text)]
    
    if clauses_in_para:
        original_text = paragraph.text
        paragraph.clear() # Erase the old text
        
        # Split the text exactly where the clauses are
        pattern = r'(\b(?:' + '|'.join(map(re.escape, clauses_in_para)) + r')\b)'
        parts = re.split(pattern, original_text)
        
        # Redraw the paragraph
        for part in parts:
            if not part: continue
            run = paragraph.add_run(part)
            # If this piece of text is our clause, paint it yellow
            if part in clauses_in_para:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def highlight_docx(input_path, output_path, found_clauses):
    """
    Opens the original Word Document, searches for the clauses, and saves 
    a newly highlighted copy.
    """
    print("  - Generating highlighted Word Document...")
    try:
        doc = docx.Document(input_path)
        
        # Check standard paragraphs
        for paragraph in doc.paragraphs:
            _apply_highlights_to_paragraph(paragraph, found_clauses)
            
        # Check paragraphs hidden inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _apply_highlights_to_paragraph(paragraph, found_clauses)
                        
        doc.save(output_path)
        print(f"  -> Highlighted DOCX saved to: {output_path}")
    except Exception as e:
        print(f"  ! Failed to highlight Word Document: {e}")

# =============================================================================
# --- Main Execution Block ---
# =============================================================================

def main():
    print("\n--- Starting Automated Compliance Matrix Generator ---")

    # Make sure our folders exist so the script doesn't crash
    os.makedirs(DATABASE_DIR, exist_ok=True)
    os.makedirs(SOLICITATIONS_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Load all the database files into memory
    master_df = load_databases(DATABASE_DIR)
    if master_df is None:
        print("Exiting due to error loading databases.")
        return

    known_clauses = master_df[CLAUSE_COL_NAME].unique().tolist()

    # Find all the PDFs and Word docs we need to process
    doc_files = [f for f in os.listdir(SOLICITATIONS_DIR) if f.lower().endswith(('.pdf', '.docx'))]

    if not doc_files:
        print(f"\nNo PDF or DOCX files found in {SOLICITATIONS_DIR}.")
        return

    # Process each document one by one
    for file_name in doc_files:
        print(f"\n--- Processing: {file_name} ---")
        file_path = os.path.join(SOLICITATIONS_DIR, file_name)

        # Get the text out of the document
        if file_name.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file_name.lower().endswith('.docx'):
            text = extract_text_from_docx(file_path)
        
        if not text:
            continue

        # Find the matching clauses
        found_clauses = find_clauses_from_db(text, known_clauses)
        print(f"Found {len(found_clauses)} unique federal clauses.")

        if not found_clauses:
            print("Warning: No matching federal clauses found in text. Skipping file save.")
            continue

        # 1. Generate and save the color-coded Excel Matrix
        compliance_matrix_df = generate_compliance_matrix(found_clauses, master_df)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        
        excel_filename = f"Compliance_Matrix_{os.path.splitext(file_name)[0]}_{timestamp}.xlsx"
        excel_path = os.path.join(OUTPUT_DIR, excel_filename)
        
        try:
            compliance_matrix_df.to_excel(excel_path, index=False)
            apply_color_coding(excel_path)
            print(f"Successfully saved formatted compliance matrix.")
        except Exception as e:
            print(f"Error saving Excel file {excel_path}: {e}")

        # 2. Generate and save the fully highlighted Document
        if file_name.lower().endswith('.pdf'):
            pdf_filename = f"Executed_Highlights_{os.path.splitext(file_name)[0]}_{timestamp}.pdf"
            pdf_output_path = os.path.join(OUTPUT_DIR, pdf_filename)
            highlight_pdf(file_path, pdf_output_path, found_clauses)
            
        elif file_name.lower().endswith('.docx'):
            docx_filename = f"Executed_Highlights_{os.path.splitext(file_name)[0]}_{timestamp}.docx"
            docx_output_path = os.path.join(OUTPUT_DIR, docx_filename)
            highlight_docx(file_path, docx_output_path, found_clauses)

    print("\n--- All tasks completed. ---")

# Standard Python command to start the script
if __name__ == '__main__':
    main()