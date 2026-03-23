"""
=============================================================================
Project: AI Contract Semantic Reviewer (Phase 2)
Description: This script reads custom/commercial contracts (Word Docs), 
             dynamically builds an AI legal playbook from the sponsor's 
             Ts&Cs Matrix, and uses an LLM to flag risky clauses and 
             suggest redlines.
=============================================================================
"""

import os
import pandas as pd
import docx
from google import genai                     # <--- NEW: Updated to the modern Google GenAI SDK
from datetime import datetime

# =============================================================================
# --- Setup & API Key ---
# =============================================================================

GEMINI_API_KEY = "API KEY HERE"

# Folder Setup
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if os.path.basename(SCRIPT_DIR).lower() == 'code':
    PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
else:
    PROJECT_DIR = SCRIPT_DIR

# The script specifically looks for the sponsor's macro-enabled matrix
MATRIX_FILE = os.path.join(PROJECT_DIR, 'Database', 'Contract Ts&Cs Matrix.xlsm')
SOLICITATIONS_DIR = os.path.join(PROJECT_DIR, 'Solicitations')
OUTPUT_DIR = os.path.join(PROJECT_DIR, 'Output')

# =============================================================================
# --- Core Functions ---
# =============================================================================

def build_ai_playbook(excel_path):
    """
    Reads the sponsor's Ts&Cs Excel Matrix and dynamically translates the 
    columns into a text-based "Playbook" that the AI can understand and enforce.
    """
    print(f"Loading Legal Playbook from: {os.path.basename(excel_path)}...")
    playbook_text = "AUBURN UNIVERSITY LEGAL PLAYBOOK\n\n"
    
    if not os.path.exists(excel_path):
        print(f"Error: Could not find the matrix file at {excel_path}")
        return ""

    try:
        xls = pd.ExcelFile(excel_path, engine='openpyxl')
        skip_sheets = ['INDEX', 'CONTACTS', 'template']
        
        # Go through every tab in the spreadsheet
        for sheet_name in xls.sheet_names:
            if sheet_name in skip_sheets:
                continue
                
            try:
                df = pd.read_excel(excel_path, sheet_name=sheet_name, engine='openpyxl')
                
                # Check for preferred language (usually row 2, column 2)
                preferred_language = str(df.iloc[1, 1]) if len(df) > 1 and len(df.columns) > 1 else 'nan'
                
                playbook_text += f"--- {sheet_name.upper()} ---\n"
                
                # Find the 'Common Problems' section
                start_row = -1
                if len(df.columns) > 1:
                    for idx, val in df.iloc[:, 1].items():
                        if str(val).strip().lower() == 'common problems':
                            start_row = idx
                            break
                
                if start_row != -1:
                    playbook_text += "UNACCEPTABLE PROVISIONS & REQUIRED RESPONSES:\n"
                    # Loop through all the problems and fallbacks listed below that row
                    for idx in range(start_row + 1, len(df)):
                        problem = str(df.iloc[idx, 1])
                        if problem == 'nan' or problem.strip() == '':
                            continue
                        
                        why = str(df.iloc[idx, 2]) if len(df.columns) > 2 else 'nan'
                        resp1 = str(df.iloc[idx, 3]) if len(df.columns) > 3 else 'nan'
                        resp2 = str(df.iloc[idx, 4]) if len(df.columns) > 4 else 'nan'
                        
                        playbook_text += f"- IF SPONSOR CONTRACT SAYS: '{problem}'\n"
                        if why != 'nan' and why.strip() != '':
                            playbook_text += f"  WHY WE REJECT THIS: {why}\n"
                        if resp1 != 'nan' and resp1.strip() != '':
                            playbook_text += f"  SUGGESTED REDLINE/RESPONSE: {resp1}\n"
                        if resp2 != 'nan' and resp2.strip() != '':
                            playbook_text += f"  FALLBACK NEGOTIATION: {resp2}\n"
                playbook_text += "\n"
            except Exception as e:
                # If a sheet is formatted weirdly, just skip it and keep going
                continue
                
        print("Playbook successfully built and loaded into AI memory.")
        return playbook_text
    except Exception as e:
        print(f"Failed to read Matrix Excel file: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Reads the raw text from the contract document."""
    print(f"Extracting text from contract: {os.path.basename(docx_path)}...")
    try:
        doc = docx.Document(docx_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip() != ""])
    except Exception as e:
        print(f"  ! Error reading {docx_path}: {e}")
        return None

def save_report_to_word(report_text, output_path):
    """Saves the AI's analysis into a clean, formatted Word document."""
    doc = docx.Document()
    doc.add_heading('Automated AI Contract Review', 0)
    
    # Simple trick to split the AI's markdown text into Word paragraphs
    for line in report_text.split('\n'):
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        elif line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=2)
        else:
            doc.add_paragraph(line)
            
    doc.save(output_path)
    print(f"  -> Saved Risk Report to: {output_path}")

def review_contract_with_ai(contract_text, playbook_text):
    """Sends the contract and the playbook to the AI model for semantic analysis."""
    print("  - Running AI Semantic Analysis (this may take 15-30 seconds)...")
    
    # NEW: Initialize the client using the new SDK syntax
    client = genai.Client(api_key=GEMINI_API_KEY)
    
    system_prompt = f"""
    You are an expert legal contract reviewer for Auburn University.
    I am going to provide you with the "AUBURN UNIVERSITY LEGAL PLAYBOOK" which outlines 
    unacceptable contract terms, the reasoning, and the exact redlines/responses you must use.
    
    Your job is to read the provided commercial contract and compare it against the Playbook.
    If you find any clauses in the contract that violate the rules in the Playbook (e.g. Governing Law outside of Alabama, unacceptable IP terms, etc.), flag them.
    
    Format your response cleanly:
    1. Quote the problematic section from the contract.
    2. Explain WHY it is an issue (using the reasoning from the Playbook).
    3. Provide the exact Suggested Redline / Response from the Playbook.
    
    Do not make up legal advice. Only flag issues that strictly violate the provided Playbook.
    
    --- PLAYBOOK ---
    {playbook_text}
    """
    
    try:
        # NEW: Call the API using the new generation method
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[system_prompt, f"--- CONTRACT TO REVIEW ---\n{contract_text}"]
        )
        return response.text
    except Exception as e:
        print(f"  ! AI API Error: {e}")
        return "An error occurred while talking to the AI API."

# =============================================================================
# --- Main Execution Block ---
# =============================================================================

def main():
    print("\n--- Starting AI Semantic Contract Reviewer ---")
    
    os.makedirs(SOLICITATIONS_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    if GEMINI_API_KEY == "YOUR_API_KEY_HERE":
        print("\n[!] ERROR: Please paste your Gemini API Key into the script!")
        return

    # 1. Build the dynamic AI Playbook from the Sponsor's Excel File
    playbook_text = build_ai_playbook(MATRIX_FILE)
    if not playbook_text:
        return

    # 2. Find any contracts that need reviewing
    doc_files = [f for f in os.listdir(SOLICITATIONS_DIR) if f.lower().endswith('.docx')]
    
    if not doc_files:
        print(f"\nNo Word Document contracts (.docx) found in {SOLICITATIONS_DIR}.")
        print("Please place a commercial contract in the folder to test.")
        return

    # 3. Process each contract
    for file_name in doc_files:
        print(f"\n--- Processing: {file_name} ---")
        file_path = os.path.join(SOLICITATIONS_DIR, file_name)
        
        contract_text = extract_text_from_docx(file_path)
        if not contract_text:
            continue
            
        # 4. Have the AI review the text against the playbook rules
        ai_report = review_contract_with_ai(contract_text, playbook_text)
        
        # 5. Save the output
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        report_filename = f"AI_Risk_Report_{os.path.splitext(file_name)[0]}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, report_filename)
        
        save_report_to_word(ai_report, output_path)

    print("\n--- All tasks completed. ---")

if __name__ == "__main__":
    main()
